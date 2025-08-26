import os
import json
import time
import pathlib
import bcrypt
import numpy as np
from numpy.linalg import norm

import streamlit as st
import google.generativeai as genai
import google.api_core.exceptions

from langchain_text_splitters import RecursiveCharacterTextSplitter
from langchain.vectorstores import FAISS
from langchain.embeddings import HuggingFaceEmbeddings
from langchain_core.documents import Document

import chardet
from PyPDF2 import PdfReader
from docx import Document as DocxDocument
from google.generativeai.types.safety_types import HarmBlockThreshold, HarmCategory

# =========================
# ---------- AUTH ----------
# =========================
USERS_PATH = pathlib.Path("users.json")

def _load_users():
    if USERS_PATH.exists():
        try:
            return json.loads(USERS_PATH.read_text(encoding="utf-8"))
        except Exception:
            pass
    # Bootstrap default admin (password: admin123) — change ASAP
    users = {"admin": {"hash": bcrypt.hashpw(b"admin123", bcrypt.gensalt()).decode("utf-8"),
                       "created_at": int(time.time())}}
    USERS_PATH.write_text(json.dumps(users, indent=2), encoding="utf-8")
    return users

def _save_users(users: dict):
    USERS_PATH.write_text(json.dumps(users, indent=2), encoding="utf-8")

def _hash_pw(pw: str) -> str:
    return bcrypt.hashpw(pw.encode("utf-8"), bcrypt.gensalt()).decode("utf-8")

def _check_pw(pw: str, hashed: str) -> bool:
    try:
        return bcrypt.checkpw(pw.encode("utf-8"), hashed.encode("utf-8"))
    except Exception:
        return False

def auth_gate():
    if "auth" not in st.session_state:
        st.session_state.auth = {"is_authed": False, "user": None}

    # After login: minimal sidebar
    if st.session_state.auth["is_authed"]:
        st.sidebar.markdown(f"✅ **Logged in as:** `{st.session_state.auth['user']}`")
        if st.sidebar.button("Logout", use_container_width=True):
            st.session_state.auth = {"is_authed": False, "user": None}
            st.rerun()
        return True, st.session_state.auth["user"]

    # Before login: show login/signup
    st.sidebar.markdown("### 🔐 Authentication")
    tab_login, tab_signup = st.sidebar.tabs(["Login", "Sign up"])

    with tab_login:
        user = st.text_input("Username", key="login_user")
        pw = st.text_input("Password", type="password", key="login_pw")
        if st.button("Login", use_container_width=True):
            users = _load_users()
            if user in users and _check_pw(pw, users[user]["hash"]):
                st.session_state.auth = {"is_authed": True, "user": user}
                st.success(f"Welcome, {user}!")
                st.rerun()
            else:
                st.error("Invalid username or password.")

    with tab_signup:
        new_user = st.text_input("New username", key="signup_user")
        new_pw = st.text_input("New password", type="password", key="signup_pw")
        new_pw2 = st.text_input("Confirm password", type="password", key="signup_pw2")
        if st.button("Create account", use_container_width=True):
            if not new_user or not new_pw:
                st.warning("Username and password are required.")
            elif len(new_pw) < 8:
                st.warning("Use at least 8 characters for the password.")
            elif new_pw != new_pw2:
                st.warning("Passwords do not match.")
            else:
                users = _load_users()
                if new_user in users:
                    st.error("Username already exists.")
                else:
                    users[new_user] = {"hash": _hash_pw(new_pw), "created_at": int(time.time())}
                    _save_users(users)
                    st.success("Account created. You can log in now.")

    return False, None

# Require login BEFORE anything else
authed, current_user = auth_gate()
if not authed:
    st.stop()

# ==============================
# ---------- SETTINGS ----------
# ==============================
st.title("📚 RAG 2.0")
st.caption(f"Logged in as **{current_user}**")

# Gemini API key from secrets
genai.configure(api_key=st.secrets.get("GEMINI_API_KEY", ""))
if not st.secrets.get("GEMINI_API_KEY"):
    st.error("Gemini API key missing. Add it to `.streamlit/secrets.toml` as GEMINI_API_KEY and restart.")
    st.stop()

# Safety settings
safety_settings = {
    HarmCategory.HARM_CATEGORY_DANGEROUS_CONTENT: HarmBlockThreshold.BLOCK_LOW_AND_ABOVE,
    HarmCategory.HARM_CATEGORY_HATE_SPEECH: HarmBlockThreshold.BLOCK_LOW_AND_ABOVE,
    HarmCategory.HARM_CATEGORY_HARASSMENT: HarmBlockThreshold.BLOCK_LOW_AND_ABOVE,
    HarmCategory.HARM_CATEGORY_SEXUALLY_EXPLICIT: HarmBlockThreshold.BLOCK_LOW_AND_ABOVE,
}

# Embeddings & per-user FAISS location
embedding_model = HuggingFaceEmbeddings(model_name="sentence-transformers/all-MiniLM-L6-v2")

def _user_index_dir(user: str) -> str:
    path = os.path.join("faiss_index", user)
    os.makedirs(path, exist_ok=True)
    return path

vector_store = None
index_dir = _user_index_dir(current_user)
if os.path.exists(index_dir):
    try:
        vector_store = FAISS.load_local(index_dir, embeddings=embedding_model, allow_dangerous_deserialization=True)
    except Exception:
        vector_store = None

# Prompt template
prompt_template = """
You are a **Safe and Helpful AI Assistant** specialized in answering user questions using only the provided context.
Rules:
- Only answer from the context below.
- If the context does not contain the answer, reply exactly: "Answer is not in the context".
- Do not fabricate information.
- Keep answers clear, concise, and well-formatted (bullet points or numbered steps).
- Professional, respectful tone. No harassment, hate, or harmful instructions.

### Context (top relevant chunks in ranked order):
{context}

### User Question:
{question}

### Reasoning Instructions:
1. Verify whether the answer exists in the provided context.
2. If yes, produce a concise answer (bullet points / numbered steps).
3. If no, reply: "Answer is not in the context".
4. Avoid filler; focus only on relevant facts.

### Final Answer:
"""

# Sidebar controls
st.sidebar.header("⚙️ Settings")
k_candidates = st.sidebar.number_input("FAISS candidates (fast retrieve)", min_value=3, max_value=50, value=10, step=1)
k_context = st.sidebar.number_input("Top chunks to include in prompt", min_value=1, max_value=10, value=5, step=1)

# ===============================
# ---------- LOAD/INDEX ----------
# ===============================
def load_and_index(file):
    file_type = file.name.split(".")[-1].lower()

    if file_type == "pdf":
        reader = PdfReader(file)
        file_content = "\n".join([page.extract_text() for page in reader.pages if page.extract_text()])
    elif file_type == "docx":
        docx_obj = DocxDocument(file)
        file_content = "\n".join([para.text for para in docx_obj.paragraphs])
    else:  # txt / md
        raw_data = file.read()
        detected_encoding = chardet.detect(raw_data)["encoding"] or "utf-8"
        try:
            file_content = raw_data.decode(detected_encoding, errors="replace")
        except UnicodeDecodeError:
            file_content = raw_data.decode("utf-8", errors="replace")

    base_doc = Document(page_content=file_content)
    text_splitter = RecursiveCharacterTextSplitter(chunk_size=1000, chunk_overlap=200)
    chunks = text_splitter.split_documents([base_doc])

    global vector_store, index_dir
    vector_store = FAISS.from_documents(chunks, embedding_model)
    vector_store.save_local(index_dir)
    st.success("✅ Document indexed successfully!")

uploaded_file = st.sidebar.file_uploader("📂 Upload a document:", type=["txt", "md", "pdf", "docx"])
if uploaded_file and st.sidebar.button("📌 Load & Index"):
    load_and_index(uploaded_file)

# =====================================
# ---------- RERANK & CONTEXT ----------
# =====================================
def cosine_similarity(a: np.ndarray, b: np.ndarray) -> float:
    if a is None or b is None:
        return -1.0
    denom = (norm(a) * norm(b))
    if denom == 0:
        return -1.0
    return float(np.dot(a, b) / denom)

def build_context_with_semantic_rerank(question: str, k_candidates: int = 10, k_context: int = 5):
    """
    1) Get top k_candidates quickly from FAISS.
    2) Compute embedding for the question and each candidate chunk.
    3) Rerank by cosine similarity and return the top k_context chunks joined as context.
    """
    if vector_store is None:
        return ""

    # 1) Candidate retrieval
    try:
        candidates = vector_store.similarity_search_with_score(question, k=k_candidates)
    except TypeError:
        candidates = [(doc, None) for doc in vector_store.similarity_search(question, k=k_candidates)]

    # 2) Embeddings
    q_emb = embedding_model.embed_query(question)
    texts = [doc.page_content for doc, _ in candidates]
    if not texts:
        return ""

    doc_embs = embedding_model.embed_documents(texts)

    # 3) Cosine rerank
    rerank_list = []
    for (doc, _), emb in zip(candidates, doc_embs):
        sim = cosine_similarity(np.array(q_emb), np.array(emb))
        rerank_list.append((sim, doc))

    rerank_list.sort(key=lambda x: x[0], reverse=True)

    # 4) Build context with scores
    selected = rerank_list[:k_context]
    context_blocks = []
    for idx, (sim, doc) in enumerate(selected, start=1):
        header = f"[Chunk {idx} — score: {sim:.4f}]"
        context_blocks.append(header + "\n" + doc.page_content.strip())

    context = "\n\n---\n\n".join(context_blocks)
    return context

# ==================================
# ---------- GENERATION ------------
# ==================================
def retrieve_and_generate(question):
    if vector_store is None:
        st.error("⚠ Please upload and index a document first.")
        return None

    context = build_context_with_semantic_rerank(
        question, k_candidates=int(k_candidates), k_context=int(k_context)
    )
    if not context:
        st.warning("No context found from indexed documents.")
        return None

    model = genai.GenerativeModel("gemini-1.5-flash-latest")

    try:
        response = model.generate_content(
            prompt_template.format(context=context, question=question),
            safety_settings=safety_settings
        )
        return response.text
    except google.api_core.exceptions.ResourceExhausted:
        st.error("❌ Gemini API quota exceeded. Please wait or upgrade your plan.")
        st.info("🔗 https://ai.google.dev/gemini-api/docs/rate-limits")
        return None
    except Exception as e:
        st.error(f"❌ Error: {e}")
        return None

# ==========================
# ---------- UI ------------
# ==========================
question = st.text_input("💬 Ask a question:")
if st.button("🤖 Get Answer"):
    if question.strip():
        with st.spinner("Searching and reranking…"):
            answer = retrieve_and_generate(question)
        if answer:
            st.markdown("### 💡 Answer:")
            st.write(answer)
    else:
        st.warning("⚠ Please enter a question before submitting.")
